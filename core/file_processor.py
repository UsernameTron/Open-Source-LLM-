import PyPDF2
try:
    from docx import Document
except ImportError:
    Document = None
from pathlib import Path
from typing import Dict, Any

class FileProcessor:
    @staticmethod
    def process_file(file_path: str) -> Dict[str, Any]:
        """Process different file types and extract text content."""
        path = Path(file_path)
        
        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
            
        try:
            if path.suffix.lower() == '.txt':
                return FileProcessor._process_txt(path)
            elif path.suffix.lower() == '.pdf':
                return FileProcessor._process_pdf(path)
            elif path.suffix.lower() in ['.doc', '.docx']:
                return FileProcessor._process_docx(path)
            else:
                raise ValueError(f"Unsupported file type: {path.suffix}")
        except Exception as e:
            raise RuntimeError(f"Error processing file {file_path}: {str(e)}")

    @staticmethod
    def _process_txt(path: Path) -> Dict[str, Any]:
        """Process text files."""
        try:
            with path.open('r', encoding='utf-8') as f:
                content = f.read()
            return {
                'content': content,
                'type': 'txt',
                'pages': 1
            }
        except UnicodeDecodeError:
            # Try with a different encoding if UTF-8 fails
            with path.open('r', encoding='latin-1') as f:
                content = f.read()
            return {
                'content': content,
                'type': 'txt',
                'pages': 1
            }

    @staticmethod
    def _process_pdf(path: Path) -> Dict[str, Any]:
        """Process PDF files."""
        with path.open('rb') as f:
            pdf_reader = PyPDF2.PdfReader(f)
            content = []
            for page in pdf_reader.pages:
                content.append(page.extract_text())
        
        return {
            'content': '\n'.join(content),
            'type': 'pdf',
            'pages': len(pdf_reader.pages)
        }

    @staticmethod
    def _process_docx(path: Path) -> Dict[str, Any]:
        """Process Word documents."""
        doc = docx.Document(path)
        content = []
        for para in doc.paragraphs:
            if para.text.strip():
                content.append(para.text)
        
        return {
            'content': '\n'.join(content),
            'type': 'docx',
            'pages': len(doc.paragraphs)  # This is an approximation
        }
